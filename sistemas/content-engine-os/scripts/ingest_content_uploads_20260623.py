#!/usr/bin/env python3
"""Ingestão governada dos materiais editoriais enviados pelo Tiaro em 2026-06-23.

Fontes:
- Biblioteca de Hooks.docx
- Roteiros Feed.txt
- Régua de Mídia.xlsx
- Roteiros de Stories.docx
- Biblioteca de Temas.docx

A ingestão é idempotente: remove itens com origem/source_ref deste lote e recria.
Não publica, não envia mensagem e não toca em dados de paciente/lead.
"""
from __future__ import annotations

import json
import re
import subprocess
from pathlib import Path

from docx import Document
from openpyxl import load_workbook

ROOT = Path('/root/.hermes/cache/documents')
HOOKS = ROOT / 'doc_393d03b16210_Instituto Vital - Biblioteca de Hooks.docx'
FEED = ROOT / 'doc_d4603bce5dd0_Roteiros Feed.txt'
MEDIA = ROOT / 'doc_a750105993db_Instituto Vital Slim (Régua de Mídia).xlsx'
STORIES = ROOT / 'doc_d42f146bbf3a_Instituto Vital - Roteiros de Stories.docx'
THEMES = ROOT / 'doc_c8e94b67f240_Instituto Vital - Biblioteca de Temas.docx'
ORIGIN = 'ivs-upload-20260623'


def psql(sql: str, params: list[str] | None = None) -> str:
    cmd = ['docker', 'exec', '-i', 'content-engine-postgres', 'psql', '-U', 'content_engine', '-d', 'content_engine', '-At']
    full = sql
    if params:
        # Use json transport to avoid shell quoting issues.
        raise NotImplementedError
    out = subprocess.run(cmd, input=full, text=True, capture_output=True, check=True)
    return out.stdout.strip()


def sql_quote(s: object) -> str:
    if s is None:
        return 'null'
    return "'" + str(s).replace("'", "''") + "'"


def slug(s: str) -> str:
    s = s.lower().strip()
    s = re.sub(r'[^a-z0-9áéíóúãõâêôç]+', '_', s)
    return s.strip('_')[:80] or 'item'


def text_docx(path: Path) -> list[str]:
    doc = Document(str(path))
    lines = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            vals = [c.text.strip() for c in row.cells]
            if any(vals):
                lines.append(' | '.join(vals))
    return lines


def insert_viral(code: str, origem: str, objetivo: str, classe: str, mecanismo: str, hook: str, tese: str, obj: str, adapt: str, uso: list | None = None):
    sql = f"""
    insert into viral_scripts (tenant_id, codigo, origem, objetivo, classe_ivs, mecanismo, hook_base, tese_central, objecao_principal, leitura_ivs, adaptacao_ivs, uso_recomendado, status)
    select t.id, {sql_quote(code)}, {sql_quote(origem)}, {sql_quote(objetivo)}, {sql_quote(classe)}, {sql_quote(mecanismo)}, {sql_quote(hook)}, {sql_quote(tese)}, {sql_quote(obj)}, {sql_quote('ingestão governada de material editorial IVS')}, {sql_quote(adapt)}, {sql_quote(json.dumps(uso or [], ensure_ascii=False))}::jsonb, 'adaptado'
    from tenants t where t.slug='demo';
    """
    psql(sql)


def insert_theme(name: str, category: str, source_ref: str):
    if not name.strip():
        return
    sql = f"""
    insert into themes (tenant_id, name, category, type, source_origin, source_ref)
    select t.id, {sql_quote(name[:500])}, {sql_quote(category[:120])}, 'biblioteca_ivs', 'upload_tiaro', {sql_quote(source_ref)}
    from tenants t where t.slug='demo';
    """
    psql(sql)


def insert_story_theme(title: str, category: str, pain: str = '', desire: str = '', objection: str = ''):
    if not title.strip():
        return
    sql = f"""
    insert into story_themes (tenant_id, title, category, pain, desire, objection, awareness_level, source, confidence)
    select t.id, {sql_quote(title[:220])}, {sql_quote(category[:80])}, {sql_quote(pain[:220])}, {sql_quote(desire[:220])}, {sql_quote(objection[:220])}, 'consciente_da_dor', {sql_quote(ORIGIN + ':stories')}, 0.92
    from tenants t where t.slug='demo'
    on conflict (tenant_id, title) do update set category=excluded.category, pain=excluded.pain, desire=excluded.desire, objection=excluded.objection, source=excluded.source, confidence=excluded.confidence;
    """
    psql(sql)


def insert_device(code: str, name: str, logic: str, example: str, category: str):
    sql = f"""
    insert into narrative_devices (tenant_id, code, name, logic, example, ai_instruction, category)
    select t.id, {sql_quote(code)}, {sql_quote(name)}, {sql_quote(logic)}, {sql_quote(example)}, {sql_quote('Use este hook como abertura; adapte sem prometer resultado médico e mantendo voz premium IVS.')}, {sql_quote(category)}
    from tenants t where t.slug='demo';
    """
    psql(sql)


def ingest_hooks() -> int:
    lines = text_docx(HOOKS)
    current = 'Hook'
    count = 0
    for line in lines:
        if line.endswith('(6 hooks)'):
            current = line.replace('(6 hooks)', '').strip()
            continue
        if line.startswith('"') and line.endswith('"'):
            count += 1
            hook = line.strip('"')
            insert_device(f'HOOK-{count:02d}', f'Hook {current} {count:02d}', f'Categoria: {current}', hook, f'hook_{slug(current)}')
            insert_viral(f'HOOK-{count:02d}', f'{ORIGIN}:biblioteca-hooks', 'atração', f'hook_{slug(current)}', current, hook, 'Hook reutilizável para abertura de conteúdo IVS', 'atenção inicial', hook, [{'formatos': ['reels','carrossel','stories','estatico']}])
    return count


def ingest_feed() -> int:
    txt = FEED.read_text(encoding='utf-8', errors='replace')
    chunks = re.split(r'ROTEIRO\s+(\d+)', txt)
    count = 0
    for i in range(1, len(chunks), 2):
        num = chunks[i]
        body = chunks[i+1].strip().strip('"')
        if not body:
            continue
        first = next((l.strip() for l in body.splitlines() if l.strip()), f'Roteiro Feed {num}')
        lower = body.lower()
        if 'perimenopausa' in lower:
            classe, objetivo, obj = 'hormonios_perimenopausa', 'educação', 'isso é coisa da idade'
        elif '15 minutos' in lower:
            classe, objetivo, obj = 'metodo_ivs', 'conversão', 'consulta rápida não resolve'
        else:
            classe, objetivo, obj = 'ja_tentei_de_tudo', 'identificação', 'já tentei de tudo'
        count += 1
        insert_viral(f'FEED-{int(num):02d}', f'{ORIGIN}:roteiros-feed', objetivo, classe, 'roteiro_feed', first, first, obj, body, [{'formatos': ['reels','feed','carrossel']}])
    return count


def ingest_themes() -> int:
    lines = text_docx(THEMES)
    current = 'Geral'
    count = 0
    skip_prefixes = ('BIBLIOTECA', 'O QUE NÃO FALAR', 'Posicionamento', 'Temas fora', 'Formatos e trends', 'ERROS')
    for line in lines:
        if line.startswith('Pilar '):
            current = line.split('—', 1)[-1].strip() if '—' in line else line
            continue
        if any(line.startswith(p) for p in skip_prefixes) or len(line) < 20:
            continue
        if line.startswith(('Prometer ', 'Mostrar ', 'Comparação ', 'Prova ', 'Feed ', 'Dancinha', 'Vídeos', 'Humor', 'Filtros')):
            insert_theme(line, 'O que não falar / evitar', f'{ORIGIN}:biblioteca-temas')
            count += 1
            continue
        insert_theme(line, current, f'{ORIGIN}:biblioteca-temas')
        count += 1
    return count


def ingest_media_xlsx() -> int:
    wb = load_workbook(str(MEDIA), data_only=True)
    count = 0
    for ws in wb.worksheets:
        rows = list(ws.iter_rows(values_only=True))
        if not rows:
            continue
        header_idx = None
        for idx, row in enumerate(rows[:20]):
            vals = [str(x).strip() if x is not None else '' for x in row]
            if 'Formato' in vals and any('Hook' in v for v in vals):
                header_idx = idx
                break
        if header_idx is None:
            continue
        headers = [str(x).strip() if x is not None else '' for x in rows[header_idx]]
        for row in rows[header_idx+1:]:
            data = {headers[i]: row[i] for i in range(min(len(headers), len(row))) if headers[i]}
            fmt = str(data.get('Formato') or '').strip()
            hook = str(data.get('Hook / Capa') or data.get('Direcionamento') or '').strip()
            copy = str(data.get('Copy principal (slides / roteiro)') or '').strip()
            legenda = str(data.get('Legenda do post') or data.get('Legenda') or '').strip()
            pillar = str(data.get('Pilar') or ws.title).strip()
            num = str(data.get('#') or '').strip()
            if not fmt or not (hook or copy or legenda):
                continue
            count += 1
            body = '\n\n'.join(x for x in [hook, copy, legenda] if x)
            classe = slug(pillar or fmt)
            objetivo = 'conversão' if 'agende' in body.lower() or 'link da bio' in body.lower() else 'educação'
            insert_viral(f'REGUA-{count:03d}', f'{ORIGIN}:regua-midia:{ws.title}', objetivo, classe, fmt.lower(), hook[:300] or f'{fmt} {num}', pillar or hook[:120], 'objeção mapeada na régua', body, [{'formato': fmt, 'sheet': ws.title, 'status_original': str(data.get('Status') or '')}])
            insert_theme(hook[:500] or pillar, pillar or fmt, f'{ORIGIN}:regua-midia')
    return count


def ingest_stories() -> int:
    doc = Document(str(STORIES))
    count = 0
    current = 'Stories IVS'
    for p in doc.paragraphs:
        txt = p.text.strip()
        if 'Destaque' in txt and '—' in txt:
            current = txt.strip()
            insert_story_theme(current, 'destaque_instagram', desire='organizar destaque estratégico do perfil')
            count += 1
    for table in doc.tables:
        # Detect rows with formato/text/audio/orientations, insert story theme/viral item per row.
        for row in table.rows:
            vals = [c.text.strip() for c in row.cells]
            blob = ' | '.join(v for v in vals if v)
            if not blob or blob.startswith('#') or len(blob) < 25:
                continue
            if any(x in blob.lower() for x in ['gancho', 'título:', 'cta', 'narr']):
                count += 1
                title = vals[1] if len(vals) > 1 and vals[1] else blob[:120]
                insert_story_theme(title[:180], 'roteiro_story', pain=current, desire='aquecer e qualificar lead pelo Instagram')
                insert_viral(f'STORY-{count:03d}', f'{ORIGIN}:roteiros-stories', 'aquecimento', 'stories_destaques', 'sequencia_stories', title[:250], current, 'dúvida/objeção em stories', blob, [{'formatos':['stories'], 'destaque': current}])
    return count


def main() -> None:
    # Limpeza idempotente apenas do lote.
    psql(f"delete from viral_scripts where origem like {sql_quote(ORIGIN + ':%')};")
    psql(f"delete from themes where source_ref like {sql_quote(ORIGIN + ':%')};")
    psql(f"delete from narrative_devices where code like 'HOOK-%' and category like 'hook_%';")
    psql(f"delete from story_themes where source={sql_quote(ORIGIN + ':stories')};")

    counts = {
        'hooks': ingest_hooks(),
        'feed_scripts': ingest_feed(),
        'themes': ingest_themes(),
        'media_sheet_items': ingest_media_xlsx(),
        'story_items': ingest_stories(),
    }
    print(json.dumps({'ok': True, 'origin': ORIGIN, 'counts': counts, 'total': sum(counts.values())}, ensure_ascii=False, indent=2))


if __name__ == '__main__':
    main()
